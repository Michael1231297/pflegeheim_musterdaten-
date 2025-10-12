from io import BytesIO
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

# === Corporate Design ===
BRAND_ROT = "#e2001A"
GRAU_DUNKEL = "#333333"


def _make_bar_image(series: pd.Series, title: str, xlabel: str, ylabel: str = "Anzahl") -> BytesIO:
    """Erstellt ein Balkendiagramm im Corporate Design."""
    counts = series.value_counts().sort_index()
    
    fig, ax = plt.subplots(figsize=(8, 5))
    
    # Balken mit AWO-Rot und abgerundeten Ecken
    bars = ax.bar(
        counts.index.astype(str), 
        counts.values,
        color=BRAND_ROT,
        alpha=0.95,
        edgecolor='none'
    )
    
    # Titel und Labels - fett und dunkel
    ax.set_title(title, fontsize=16, fontweight='bold', color=GRAU_DUNKEL, pad=20)
    ax.set_xlabel(xlabel, fontsize=13, fontweight='bold', color=GRAU_DUNKEL, labelpad=10)
    ax.set_ylabel(ylabel, fontsize=13, fontweight='bold', color=GRAU_DUNKEL, labelpad=10)
    
    # Y-Achse: Nur ganze Zahlen
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    
    # Achsen-Styling - starke Kontraste
    ax.spines['bottom'].set_color(GRAU_DUNKEL)
    ax.spines['bottom'].set_linewidth(2)
    ax.spines['left'].set_color(GRAU_DUNKEL)
    ax.spines['left'].set_linewidth(2)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Tick-Styling
    ax.tick_params(axis='both', which='major', labelsize=11, width=2, color=GRAU_DUNKEL, labelcolor=GRAU_DUNKEL)
    
    # Grid mit besserer Sichtbarkeit
    ax.yaxis.grid(True, linestyle='-', alpha=0.5, color='#cccccc', linewidth=1)
    ax.set_axisbelow(True)
    
    # Werte über den Balken
    for bar, v in zip(bars, counts.values):
        height = bar.get_height()
        ax.annotate(
            f"{int(v)}",
            xy=(bar.get_x() + bar.get_width() / 2, height),
            xytext=(0, 5),
            textcoords="offset points",
            ha="center",
            va="bottom",
            fontsize=11,
            fontweight='bold',
            color=GRAU_DUNKEL
        )
    
    # X-Achsen-Labels gerade
    plt.xticks(rotation=0, ha="center")
    
    fig.tight_layout()
    
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor='white')
    plt.close(fig)
    buf.seek(0)
    
    return buf


def _make_age_group_image(df: pd.DataFrame) -> BytesIO:
    """Erstellt Altersgruppen-Diagramm (70-74, 75-79, etc.)."""
    df_age = df.copy()
    bins = [70, 75, 80, 85, 90, 95, 100]
    labels = ["70-74", "75-79", "80-84", "85-89", "90-94", "95+"]
    df_age["Altersgruppe"] = pd.cut(df_age["Alter"], bins=bins, labels=labels, right=False)
    
    # Series für _make_bar_image erstellen
    age_series = df_age["Altersgruppe"].dropna()
    
    return _make_bar_image(age_series, "Altersverteilung", "Altersgruppe", "Anzahl Bewohner")


def build_word_report(df: pd.DataFrame) -> BytesIO:
    """Erzeugt einen Word-Report mit den Grafiken im Corporate Design."""
    doc = Document()
    
    # === Titel mit Corporate Design ===
    title = doc.add_heading("Pflegeheim – Datenanalyse", 0)
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(226, 0, 26)  # AWO-Rot
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # === Einleitung ===
    intro = doc.add_paragraph(
        "Automatisch generierter Bericht aus der hochgeladenen Excel-Datei. "
        "Die folgenden Abbildungen zeigen die wichtigsten Verteilungen."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Leerzeile
    
    # === KPI-Übersicht ===
    doc.add_heading("📊 Kennzahlen im Überblick", level=1)
    
    kpi_text = f"• Bewohner gesamt: {len(df)}\n"
    
    if "Alter" in df.columns:
        durchschnittsalter = df["Alter"].mean()
        kpi_text += f"• Durchschnittsalter: {durchschnittsalter:.1f} Jahre\n"
    
    if "Betreuungsbedarf" in df.columns:
        hoher_bedarf = len(df[df["Betreuungsbedarf"] == "hoch"])
        anteil = (hoher_bedarf / len(df) * 100) if len(df) > 0 else 0
        kpi_text += f"• Hoher Betreuungsbedarf: {hoher_bedarf} ({anteil:.1f}%)\n"
    
    if "Einzelzimmer" in df.columns:
        einzelzimmer = len(df[df["Einzelzimmer"] == "Ja"])
        anteil_ez = (einzelzimmer / len(df) * 100) if len(df) > 0 else 0
        kpi_text += f"• Einzelzimmer: {einzelzimmer} ({anteil_ez:.1f}%)\n"
    
    kpi_paragraph = doc.add_paragraph(kpi_text)
    kpi_paragraph.style = 'List Bullet'
    
    doc.add_page_break()
    
    # === Charts ===
    doc.add_heading("📈 Detaillierte Auswertungen", level=1)
    
    # Altersverteilung (gruppiert)
    if "Alter" in df.columns and not df["Alter"].empty:
        doc.add_heading("Altersverteilung", level=2)
        img = _make_age_group_image(df)
        doc.add_picture(img, width=Inches(6.5))
        doc.add_paragraph()  # Leerzeile
    
    # Betreuungsbedarf
    if "Betreuungsbedarf" in df.columns and not df["Betreuungsbedarf"].empty:
        doc.add_heading("Betreuungsbedarf", level=2)
        img = _make_bar_image(df["Betreuungsbedarf"], "Verteilung Betreuungsbedarf", "Betreuungsbedarf", "Anzahl")
        doc.add_picture(img, width=Inches(6.5))
        doc.add_paragraph()  # Leerzeile
    
    # Abteilungen
    if "Abteilung" in df.columns and not df["Abteilung"].empty:
        doc.add_heading("Abteilungen", level=2)
        img = _make_bar_image(df["Abteilung"], "Verteilung nach Abteilungen", "Abteilung", "Anzahl")
        doc.add_picture(img, width=Inches(6.5))
        doc.add_paragraph()  # Leerzeile
    
    # Speichern
    mem = BytesIO()
    doc.save(mem)
    mem.seek(0)
    
    return mem
