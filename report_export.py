from io import BytesIO
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

# === Corporate Design ===
BRAND_ROT = "#e2001A"
GRAU_DUNKEL = "#333333"

# === Diagramm-Einstellungen (hier kannst du die Größe anpassen) ===
DIAGRAMM_BREITE_INCHES = 3.8  # Breite der Diagramme in Inches (Standard: 5.0)


def _make_bar_image(series: pd.Series, title: str, xlabel: str, ylabel: str = "Anzahl") -> BytesIO:
    """Erstellt ein Balkendiagramm im Corporate Design."""
    counts = series.value_counts().sort_index()
    
    fig, ax = plt.subplots(figsize=(8, 5))
    
    # Balken mit AWO-Rot
    bars = ax.bar(
        counts.index.astype(str), 
        counts.values,
        color=BRAND_ROT,
        alpha=0.95,
        edgecolor='none'
    )
    
    # Titel und Labels - fett und dunkel
    ax.set_title(title, fontsize=16, fontweight='bold', color=GRAU_DUNKEL, pad=20)
    ax.set_xlabel(xlabel, fontsize=13, fontweight='bold', color=GRAU_DUNKEL, labelpad=10)
    ax.set_ylabel(ylabel, fontsize=13, fontweight='bold', color=GRAU_DUNKEL, labelpad=10)
    
    # Y-Achse: Nur ganze Zahlen
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    
    # Achsen-Styling - starke Kontraste
    ax.spines['bottom'].set_color(GRAU_DUNKEL)
    ax.spines['bottom'].set_linewidth(2)
    ax.spines['left'].set_color(GRAU_DUNKEL)
    ax.spines['left'].set_linewidth(2)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Tick-Styling
    ax.tick_params(axis='both', which='major', labelsize=11, width=2, color=GRAU_DUNKEL, labelcolor=GRAU_DUNKEL)
    
    # Grid mit besserer Sichtbarkeit
    ax.yaxis.grid(True, linestyle='-', alpha=0.5, color='#cccccc', linewidth=1)
    ax.set_axisbelow(True)
    
    # X-Achsen-Labels gerade
    plt.xticks(rotation=0, ha="center")
    
    fig.tight_layout()
    
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor='white')
    plt.close(fig)
    buf.seek(0)
    
    return buf


def _make_age_group_image(df: pd.DataFrame) -> BytesIO:
    """Erstellt Altersgruppen-Diagramm (70-74, 75-79, etc.)."""
    df_age = df.copy()
    bins = [70, 75, 80, 85, 90, 95, 100]
    labels = ["70-74", "75-79", "80-84", "85-89", "90-94", "95+"]
    df_age["Altersgruppe"] = pd.cut(df_age["Alter"], bins=bins, labels=labels, right=False)
    
    # Series für _make_bar_image erstellen
    age_series = df_age["Altersgruppe"].dropna()
    
    return _make_bar_image(age_series, "Altersverteilung", "Altersgruppe", "Anzahl Bewohner")


def _analyze_age_distribution(df: pd.DataFrame) -> str:
    """Erstellt intelligente Analyse der Altersverteilung."""
    df_age = df.copy()
    bins = [70, 75, 80, 85, 90, 95, 100]
    labels = ["70-74", "75-79", "80-84", "85-89", "90-94", "95+"]
    df_age["Altersgruppe"] = pd.cut(df_age["Alter"], bins=bins, labels=labels, right=False)
    
    counts = df_age["Altersgruppe"].value_counts()
    total = len(df_age)
    durchschnitt = df["Alter"].mean()
    median = df["Alter"].median()
    
    # Größte Gruppe
    groesste_gruppe = counts.idxmax()
    groesste_anzahl = counts.max()
    groesste_prozent = (groesste_anzahl / total * 100)
    
    # Zweitgrößte Gruppe
    counts_sorted = counts.sort_values(ascending=False)
    zweitgroesste_gruppe = counts_sorted.index[1] if len(counts_sorted) > 1 else None
    zweitgroesste_anzahl = counts_sorted.iloc[1] if len(counts_sorted) > 1 else 0
    
    # Hochbetagte (90+)
    hochbetagte = len(df_age[df_age["Alter"] >= 90])
    hochbetagte_prozent = (hochbetagte / total * 100) if total > 0 else 0
    
    text = (
        f"Die am stärksten vertretene Altersgruppe ist {groesste_gruppe} Jahre mit {groesste_anzahl} Bewohnern "
        f"({groesste_prozent:.1f}% aller Bewohner). "
    )
    
    if zweitgroesste_gruppe:
        text += (
            f"Darauf folgt die Altersgruppe {zweitgroesste_gruppe} Jahre mit {zweitgroesste_anzahl} Bewohnern. "
        )
    
    text += (
        f"Das Durchschnittsalter beträgt {durchschnitt:.1f} Jahre, der Median liegt bei {median:.1f} Jahren. "
    )
    
    if hochbetagte_prozent > 20:
        text += (
            f"Mit {hochbetagte} Bewohnern über 90 Jahren ({hochbetagte_prozent:.1f}%) zeigt sich ein hoher Anteil "
            f"hochbetagter Personen, was besondere Anforderungen an die Pflege stellt."
        )
    elif hochbetagte > 0:
        text += (
            f"Insgesamt sind {hochbetagte} Bewohner über 90 Jahre alt ({hochbetagte_prozent:.1f}%)."
        )
    
    return text


def _analyze_betreuungsbedarf(df: pd.DataFrame) -> str:
    """Erstellt intelligente Analyse des Betreuungsbedarfs."""
    counts = df["Betreuungsbedarf"].value_counts()
    total = len(df)
    
    hoch = counts.get("hoch", 0)
    mittel = counts.get("mittel", 0)
    niedrig = counts.get("niedrig", 0)
    
    hoch_prozent = (hoch / total * 100) if total > 0 else 0
    mittel_prozent = (mittel / total * 100) if total > 0 else 0
    niedrig_prozent = (niedrig / total * 100) if total > 0 else 0
    
    # Dominante Kategorie
    dominante_kategorie = counts.idxmax()
    dominante_anzahl = counts.max()
    dominante_prozent = (dominante_anzahl / total * 100)
    
    text = (
        f"Der Betreuungsbedarf verteilt sich wie folgt: {hoch} Bewohner ({hoch_prozent:.1f}%) benötigen "
        f"einen hohen Betreuungsaufwand, {mittel} Bewohner ({mittel_prozent:.1f}%) haben einen mittleren Bedarf "
        f"und {niedrig} Bewohner ({niedrig_prozent:.1f}%) weisen einen niedrigen Betreuungsbedarf auf. "
    )
    
    if dominante_kategorie == "hoch":
        text += (
            f"Mit {dominante_prozent:.1f}% liegt der Schwerpunkt auf Bewohnern mit hohem Betreuungsbedarf, "
            f"was einen entsprechend hohen Personaleinsatz erfordert."
        )
    elif dominante_kategorie == "mittel":
        text += (
            f"Die Mehrheit der Bewohner ({dominante_prozent:.1f}%) weist einen mittleren Betreuungsbedarf auf, "
            f"was eine ausgewogene Personalplanung ermöglicht."
        )
    else:
        text += (
            f"Mit {dominante_prozent:.1f}% sind die meisten Bewohner weitgehend selbstständig, "
            f"was die Pflegeintensität insgesamt reduziert."
        )
    
    return text


def _analyze_abteilungen(df: pd.DataFrame) -> str:
    """Erstellt intelligente Analyse der Abteilungsverteilung."""
    counts = df["Abteilung"].value_counts()
    total = len(df)
    
    # Größte Abteilung
    groesste_abt = counts.idxmax()
    groesste_anzahl = counts.max()
    groesste_prozent = (groesste_anzahl / total * 100)
    
    # Kleinste Abteilung
    kleinste_abt = counts.idxmin()
    kleinste_anzahl = counts.min()
    kleinste_prozent = (kleinste_anzahl / total * 100)
    
    # Auslastungsanalyse
    anzahl_abteilungen = len(counts)
    durchschnitt_pro_abt = total / anzahl_abteilungen
    
    text = (
        f"Die Bewohner verteilen sich auf {anzahl_abteilungen} Abteilungen. "
        f"Die {groesste_abt} ist mit {groesste_anzahl} Bewohnern ({groesste_prozent:.1f}%) am stärksten belegt. "
    )
    
    text += (
        f"Die {kleinste_abt} weist mit {kleinste_anzahl} Bewohnern ({kleinste_prozent:.1f}%) die geringste Belegung auf. "
    )
    
    # Gleichverteilung prüfen
    diff_prozent = groesste_prozent - kleinste_prozent
    if diff_prozent < 15:
        text += (
            f"Die Verteilung ist mit durchschnittlich {durchschnitt_pro_abt:.1f} Bewohnern pro Abteilung "
            f"relativ ausgeglichen, was eine gleichmäßige Ressourcenverteilung begünstigt."
        )
    else:
        text += (
            f"Die Unterschiede in der Belegung sind mit einer Differenz von {diff_prozent:.1f} Prozentpunkten "
            f"deutlich ausgeprägt. Dies sollte bei der Personalplanung berücksichtigt werden."
        )
    
    return text


def build_word_report(df: pd.DataFrame) -> BytesIO:
    """Erzeugt einen Word-Report mit den Grafiken im Corporate Design."""
    doc = Document()
    
    # === Titel mit Corporate Design ===
    title = doc.add_heading("Pflegeheim – Datenanalyse", 0)
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(226, 0, 26)  # BRAND-Rot
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # === Einleitung ===
    intro = doc.add_paragraph(
        "Dieser Bericht wurde automatisch auf Basis der hochgeladenen Excel-Datei erstellt und bietet "
        "eine umfassende Analyse der aktuellen Bewohnerstruktur. Die folgenden Auswertungen geben Einblicke "
        "in Altersverteilung, Betreuungsbedarf und Abteilungsbelegung."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Leerzeile
    
    # === KPI-Übersicht ===
    heading_kpi = doc.add_heading("📊 Kennzahlen im Überblick", level=1)
    for run in heading_kpi.runs:
        run.font.color.rgb = RGBColor(226, 0, 26)
    
    kpi_lines = [f"Bewohner gesamt: {len(df)}"]
    
    if "Alter" in df.columns:
        durchschnittsalter = df["Alter"].mean()
        kpi_lines.append(f"Durchschnittsalter: {durchschnittsalter:.1f} Jahre")
    
    if "Betreuungsbedarf" in df.columns:
        hoher_bedarf = len(df[df["Betreuungsbedarf"] == "hoch"])
        anteil = (hoher_bedarf / len(df) * 100) if len(df) > 0 else 0
        kpi_lines.append(f"Hoher Betreuungsbedarf: {hoher_bedarf} ({anteil:.1f}%)")
    
    if "Einzelzimmer" in df.columns:
        einzelzimmer = len(df[df["Einzelzimmer"] == "Ja"])
        anteil_ez = (einzelzimmer / len(df) * 100) if len(df) > 0 else 0
        kpi_lines.append(f"Einzelzimmer: {einzelzimmer} ({anteil_ez:.1f}%)")
    
    for line in kpi_lines:
        p = doc.add_paragraph(line, style='List Bullet')
    
    doc.add_paragraph()  # Leerzeile
    
    # === Charts ===
    heading_charts = doc.add_heading("📈 Detaillierte Auswertungen", level=1)
    for run in heading_charts.runs:
        run.font.color.rgb = RGBColor(226, 0, 26)
    
    # Altersverteilung (gruppiert)
    if "Alter" in df.columns and not df["Alter"].empty:
        heading_age = doc.add_heading("Altersverteilung", level=2)
        for run in heading_age.runs:
            run.font.color.rgb = RGBColor(226, 0, 26)
        
        # Analyse-Text
        analyse_text = _analyze_age_distribution(df)
        doc.add_paragraph(analyse_text)
        doc.add_paragraph()  # Leerzeile
        
        # Diagramm
        img = _make_age_group_image(df)
        doc.add_picture(img, width=Inches(DIAGRAMM_BREITE_INCHES))
        doc.add_paragraph()  # Leerzeile
    
    # Betreuungsbedarf
    if "Betreuungsbedarf" in df.columns and not df["Betreuungsbedarf"].empty:
        heading_bedarf = doc.add_heading("Betreuungsbedarf", level=2)
        for run in heading_bedarf.runs:
            run.font.color.rgb = RGBColor(226, 0, 26)
        
        # Analyse-Text
        analyse_text = _analyze_betreuungsbedarf(df)
        doc.add_paragraph(analyse_text)
        doc.add_paragraph()  # Leerzeile
        
        # Diagramm
        img = _make_bar_image(df["Betreuungsbedarf"], "Verteilung Betreuungsbedarf", "Betreuungsbedarf", "Anzahl")
        doc.add_picture(img, width=Inches(DIAGRAMM_BREITE_INCHES))
        doc.add_paragraph()  # Leerzeile
    
    # Abteilungen
    if "Abteilung" in df.columns and not df["Abteilung"].empty:
        heading_abt = doc.add_heading("Abteilungen", level=2)
        for run in heading_abt.runs:
            run.font.color.rgb = RGBColor(226, 0, 26)
        
        # Analyse-Text
        analyse_text = _analyze_abteilungen(df)
        doc.add_paragraph(analyse_text)
        doc.add_paragraph()  # Leerzeile
        
        # Diagramm
        img = _make_bar_image(df["Abteilung"], "Verteilung nach Abteilungen", "Abteilung", "Anzahl")
        doc.add_picture(img, width=Inches(DIAGRAMM_BREITE_INCHES))
        doc.add_paragraph()  # Leerzeile
    
    # Speichern
    mem = BytesIO()
    doc.save(mem)
    mem.seek(0)
    
    return mem


